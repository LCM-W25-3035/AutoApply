# Reference
# (OpenAI, ChatGPT o1, first prompt, 2025): I have this template in Word, this Json, both have the same keys, guide me to make a code that reeplace the information in the template with the info in Json
# (Claude, 3.5 Sonnet, last prompt, 2025): The template is not filling completly good, help me to correct the format

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json

class CVGenerator:
    CUSTOM_BULLET_STYLE = 'Custom Bullet'
    
    def __init__(self, template_path):
        self.doc = Document(template_path)
        self.setup_styles()
        
    def setup_styles(self):
        """Configure document styles"""
        styles = self.doc.styles
        
        # Normal style
        style = styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        
        # Custom bullet style
        if self.CUSTOM_BULLET_STYLE not in styles:
            bullet_style = styles.add_style(self.CUSTOM_BULLET_STYLE, WD_STYLE_TYPE.PARAGRAPH)
            bullet_style.base_style = styles['Normal']
            bullet_style.paragraph_format.left_indent = Inches(0.25)
            bullet_style.paragraph_format.first_line_indent = Inches(-0.25)
            
    def add_bullet_paragraph(self, text):
        """Add a bullet point paragraph"""
        paragraph = self.doc.add_paragraph()
        paragraph.style = self.CUSTOM_BULLET_STYLE
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.add_run('• ').bold = True
        paragraph.add_run(text.strip())
        return paragraph
        
    def add_section_title(self, title):
        """Add a section title with proper formatting"""
        # Add space before section
        spacing_para = self.doc.add_paragraph()
        spacing_para.paragraph_format.space_before = Pt(12)
        
        # Add section title
        paragraph = self.doc.add_paragraph()
        paragraph.add_run(title).bold = True
        self.add_horizontal_line(paragraph)
        
    def add_horizontal_line(self, paragraph):
        """Add horizontal line after section titles"""
        p = paragraph._p
        p_pr = p.get_or_add_pPr()
        bottom_border = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom_border.append(bottom)
        p_pr.append(bottom_border)
        
    def add_right_aligned_text(self, paragraph, text):
        """Add right-aligned text to a paragraph using tab stops"""
        # Clear any existing tab stops
        paragraph.paragraph_format.tab_stops.clear_all()
        
        # Add a right-aligned tab stop at 6 inches
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(6), WD_TAB_ALIGNMENT.RIGHT
        )
        
        # Add the text with tab
        paragraph.add_run('\t' + text)
        
    def fill_cv(self, data):
        """Fill the CV with the provided data"""
        # Clear the document
        for paragraph in self.doc.paragraphs[:]:
            p = paragraph._element
            p.getparent().remove(p)
            
        # Name
        name_paragraph = self.doc.add_paragraph()
        name_run = name_paragraph.add_run(data['PersonalInfo']['Name'])
        name_run.bold = True
        name_run.font.size = Pt(48)
        
        # Contact Info
        contact = f"{data['PersonalInfo']['Address']} | {data['PersonalInfo']['Phone']} | {data['PersonalInfo']['Email']}"
        contact_paragraph = self.doc.add_paragraph()
        contact_paragraph.add_run(contact)
        contact_paragraph.paragraph_format.space_after = Pt(12)
        
        # Summary
        self.add_section_title("Summary")
        summary_para = self.doc.add_paragraph()
        summary_para.add_run(data['Summary'])
        summary_para.paragraph_format.space_after = Pt(12)
        
        # Experience
        self.add_section_title("Experience")
        for exp in sorted(data['Experience'], key=lambda x: x['Dates'].split('-')[1], reverse=True):
            # Company and dates
            exp_para = self.doc.add_paragraph()
            exp_para.paragraph_format.space_before = Pt(6)
            company_run = exp_para.add_run(f"{exp['Company']}")
            company_run.bold = True
            self.add_right_aligned_text(exp_para, exp['Dates'])
            
            # Functions
            functions = exp['Functions'].split('\n')
            for function in functions:
                if function.strip():
                    bullet_para = self.add_bullet_paragraph(function.strip())
                    bullet_para.paragraph_format.space_after = Pt(3)
        
        # Education
        self.add_section_title("Education")
        for edu in sorted(data['Education'], key=lambda x: x['Dates'].split('-')[1], reverse=True):
            # Institution and dates
            edu_para = self.doc.add_paragraph()
            edu_para.paragraph_format.space_before = Pt(6)
            institution_run = edu_para.add_run(f"{edu['Institution']}")
            institution_run.bold = True
            self.add_right_aligned_text(edu_para, edu['Dates'])
            
            # Degree on next line
            degree_para = self.doc.add_paragraph()
            degree_para.add_run(edu['Degree'])
            degree_para.paragraph_format.space_after = Pt(6)
        
        # Skills
        self.add_section_title("Skills")
        skills = data['Skills'].split(', ')
        for skill in skills:
            bullet_para = self.add_bullet_paragraph(skill.strip())
            bullet_para.paragraph_format.space_after = Pt(3)
    
    def save(self, output_path):
        """Save the document to the specified path"""
        self.doc.save(output_path)

def generate_cv(json_file_path, template_path, output_path):
    """Generate CV from JSON data"""
    try:
        # Load JSON data
        with open(json_file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
        
        # Generate CV
        generator = CVGenerator(template_path)
        generator.fill_cv(data)
        generator.save(output_path)
        
        print(f"CV generated successfully: {output_path}")
        return True
        
    except Exception as e:
        print(f"Error generating CV: {str(e)}")
        return False

if __name__ == "__main__":
    # Example usage
    generate_cv(
        'customized_cv_output.json',
        'template1.docx',
        'cv_filled.docx'
    )
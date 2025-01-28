import openai
from docx import Document
import requests
import logging

# Basic logging configuration
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Ollama API configuration
openai.api_base = "http://localhost:11434/v1"
openai.api_key = "ollama"

def load_cv(file_path):
    """Loads the content of a CV from a Word file."""
    try:
        doc = Document(file_path)
        content = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        return content
    except Exception as e:
        logging.error(f"Error loading the file {file_path}: {e}")
        return ""

def save_cv(content, file_path):
    """Saves the customized CV content into a Word file."""
    try:
        doc = Document()
        for line in content.split("\n"):
            if line.strip():  # Ignore empty lines
                doc.add_paragraph(line)
        doc.save(file_path)
        logging.info(f"Customized CV saved at: {file_path}")
    except Exception as e:
        logging.error(f"Error saving the file {file_path}: {e}")

def analyze_cv_skills(cv_content):
    """Analyzes the CV and extracts only the key skills."""
    prompt = f"""
    Analyze the following CV and extract:
    1. A clear and complete list of all the skills mentioned in the document.

    CV:
    {cv_content}
    """
    return send_prompt_to_ollama(prompt)

def analyze_job_requirements(job_description):
    """Analyzes the job description and extracts key requirements."""
    prompt = f"""
    Analyze the following job description and extract:
    1. The key requirements (skills, experience, tools).
    2. The primary responsibilities of the role.

    Job description:
    {job_description}
    """
    return send_prompt_to_ollama(prompt)

def generate_cv_recommendations(cv_skills, job_requirements):
    """Compares the CV analysis and job description to generate recommendations."""
    prompt = f"""
    Based on the following analyses:
    1. CV skills: {cv_skills}
    2. Job description requirements: {job_requirements}

    Generate detailed recommendations to customize the CV. Identify:
    - Additional skills or experiences that should be highlighted.
    - Areas that need adjustments or emphasis to better align with the job description.
    """
    return send_prompt_to_ollama(prompt)

def send_prompt_to_ollama(prompt):
    """Sends a prompt to Ollama and returns the response."""
    url = "http://localhost:11434/v1/completions"
    payload = {
        "model": "llama3.1:8b",
        "prompt": prompt,
        "max_tokens": 1500
    }

    try:
        logging.info("Sending the following prompt to the server:")
        logging.info(prompt)  # Logs the prompt for debugging
        response = requests.post(url, json=payload, timeout=60)
        response.raise_for_status()
        return response.json().get("text", "").strip()
    except requests.exceptions.HTTPError as http_err:
        logging.error(f"HTTP error: {http_err}")
        if response is not None:
            logging.error(f"Server response: {response.text}")
        return ""
    except requests.exceptions.RequestException as req_err:
        logging.error(f"Request error to Ollama: {req_err}")
        return ""

if __name__ == "__main__":
    # File paths and job description
    cv_path = "C:/Users/diana/OneDrive/Escritorio/Test/Diana Mayorga CV.docx"
    output_path = "C:/Users/diana/OneDrive/Escritorio/Test/customized_cv.docx"
    job_description = """We are seeking a skilled and passionate Data Analytics Professional to join our dynamic team at Quarter4 Inc. As a key player in the iGaming industry, we are committed to leveraging cutting-edge technology and data analytics to enhance our sports projections and provide our customers with the most accurate insights.

Key responsibilities:
Collaborate closely with our Machine Learning, AI and Product teams to analyze historical sports data and identify trends, patterns, and valuable features.
Develop and implement advanced analytics techniques to improve the accuracy of our sports projections.
Perform thorough analysis and reporting on our projections, providing insights and recommendations to optimize our betting strategies.
Work cross-functionally with other teams to integrate data analytics into various aspects of our business operations.
Stay up-to-date with the latest developments in data analytics, machine learning, and sports betting to continuously improve our methodologies.

Your Skillset:
Bachelor's or Master's degree in Computer Science, Statistics, Mathematics, or a related field.
3 to 5+ years of proven experience in data analytics, preferably in the iGaming or related industry.
Experience analyzing large data sets.
Proficiency in programming languages such as Python, and MySQL.
Strong analytical and problem-solving skills, with the ability to interpret complex data sets and draw meaningful insights.
Experience with data visualization tools such as Tableau or Power BI is a must.
Excellent communication and collaboration skills, with the ability to effectively communicate technical concepts to non-technical stakeholders.
"""

    # Load the CV
    cv_content = load_cv(cv_path)
    if not cv_content:
        print("Failed to load the CV.")
        exit()

    # Step 1: Analyze CV skills
    logging.info("Analyzing CV skills...")
    cv_skills = analyze_cv_skills(cv_content)
    if not cv_skills:
        print("Error analyzing CV skills.")
        exit()
    print("Extracted skills from the CV:")
    print(cv_skills)

    # Step 2: Analyze job requirements
    logging.info("Analyzing job description requirements...")
    job_requirements = analyze_job_requirements(job_description)
    if not job_requirements:
        print("Error analyzing job description.")
        exit()
    print("Key requirements from the job description:")
    print(job_requirements)

    # Step 3: Generate recommendations
    logging.info("Generating recommendations to customize the CV...")
    recommendations = generate_cv_recommendations(cv_skills, job_requirements)
    if not recommendations:
        print("Error generating recommendations.")
        exit()
    print("Recommendations to customize the CV:")
    print(recommendations)

    # Confirm and save the customized CV
    user_response = input("Do you want to generate the customized CV with these recommendations? (y/n): ").strip().lower()
    if user_response == "y":
        save_cv(recommendations, output_path)
    else:
        print("The customized CV was not generated.")


import json
import streamlit as st
import pandas as pd
import google.generativeai as genai
from pypdf import PdfReader
from io import BytesIO
import re
import string
import time

your_api_key = st.secrets["api_keys"]["GEMINI_API_KEY"]
model_gemini = "models/gemini-2.0-flash"
clean_json = "```json\n"
var_resume_json = "resume/resume.json"
var_fill_cv = "fill_cv line"
var_resume_match_skils = "resume/resume_match_skills.json"

def extract_keywords_with_gemini():

    with open(var_resume_json, "r", encoding="utf-8") as file:
        resume_json = json.load(file)

    validation_prompt = f"""You are a resume keyword extractor.

Given the following resume content, extract up to 25 relevant keywords that best represent this candidate's profile. Each keyword should be 1 to 3 words long, and focus on:
- Technical skills
- Tools and technologies
- Domain expertise
- Roles and methodologies (e.g. MLOps, Data Engineering)

Resume content:
{json.dumps(resume_json)}

Return only a Python list of keywords, no explanation.
"""
    model = genai.GenerativeModel(
        model_gemini,
        generation_config={
            "temperature": 0.0,
            "top_p": 1.0,
            "top_k": 1,
            "max_output_tokens": 1024
        }
    )

    try:
        response = model.generate_content(validation_prompt)
        keywords = response.text.strip()
        if keywords.startswith("```"):
            keywords = keywords.strip("```python").strip("```").strip()
        keywords = eval(keywords)
        if isinstance(keywords, list):
            return keywords
    except Exception as e:
        print(f"❌ Error extracting keywords from Gemini: {e}")
    
    return []


def key_words_match_jobs_resume(
    f_key_words_candidate,
    f_df_text,
    max_retries=2
):
    genai.configure(api_key=your_api_key)

    validation_prompt = f"""
    You are an AI job matching assistant.

    Given a candidate with the following keywords:
    {f_key_words_candidate}

    And the following job postings:
    {f_df_text}

    For each job, return:
    - The `Job ID` of the job
    - The number of keywords from `key_words_app` that match the candidate's skills (`matches`)
    - The total number of keywords in the job's `key_words_app` (`total_of_keywords`)
    - The Jaccard similarity between the candidate's keywords and the job's `key_words_app`, calculated as:
  (number of overlapping keywords) / (total number of unique keywords between both sets)* 100,
  rounded to two decimal places. Name this field `similarity`

    Be flexible. Accept close matches or inferred relationships (e.g., "Cloud Computing" ↔ "AWS", or "ETL" ↔ "Spoon").

    Only return the **top 10 job postings** sorted in **descending order by percentage**.

    🧾 Format your answer strictly as a Python list of dictionaries like this:
    [
    {{'Job ID': 'abc123', 'matches': 7, 'total_of_keywords': 10, 'similarity': 70.0}},
    {{'Job ID': 'abc456', 'matches': 3, 'total_of_keywords': 6, 'similarity': 50.0}}
    ]

    ❌ Do not include any extra fields, explanations, or markdown formatting.
    Return only the raw Python list.
    """

    model = genai.GenerativeModel(
        model_gemini,
        generation_config={
            "temperature": 0.0,
            "top_p": 1.0,
            "top_k": 1,
            "max_output_tokens": 1024
        }
    )

    attempt = 0
    while attempt < max_retries:
        try:
            response = model.generate_content(validation_prompt)
            result_text = response.text.strip()
            if result_text.startswith("```"):
                result_text = result_text.strip("```python").strip("```").strip()
            
            result_data = eval(result_text)
            if isinstance(result_data, list) and all(
                isinstance(item, dict) and
                "Job ID" in item and
                "matches" in item and
                "total_of_keywords" in item
                for item in result_data
            ):
                return result_data

            else:
                print(f"⚠️ Attempt {attempt + 1}: Invalid format. Retrying...")
                attempt += 1
                time.sleep(1)

        except Exception as e:
            print(f"❌ Error on attempt {attempt + 1}: {e}")
            attempt += 1
            time.sleep(1)

    return False, "❌ Gemini did not return a valid structured response after multiple attempts."

def export_match_and_missing_skills():

    ats_result_path="resume/ats_score_evaluation_pre.json"
    with open(ats_result_path, "r", encoding="utf-8") as file:
        ats_result = json.load(file)

    match_skills = {
        "technical_skills": ats_result.get("matching_technical_skills", []),
        "soft_skills": ats_result.get("matching_soft_skills", [])
    }

    missing_skills = {
        "technical_skills": ats_result.get("missing_technical_skills", []),
        "soft_skills": ats_result.get("missing_soft_skills", [])
    }

    match_path = var_resume_match_skils
    missing_path = "resume/resume_missing_skills.json"

    with open(match_path, "w", encoding="utf-8") as file:
        json.dump(match_skills, file, indent=4, ensure_ascii=False)
        print(f"✅ Saved: {match_path}")

    with open(missing_path, "w", encoding="utf-8") as file:
        json.dump(missing_skills, file, indent=4, ensure_ascii=False)
        print(f"✅ Saved: {missing_path}")


def ats_score_evaluation_pre():

    with open(var_resume_json, "r", encoding="utf-8") as file:
        resume = json.load(file)

    with open("resume/job_posting.json", "r", encoding="utf-8") as file:
        job_posting = json.load(file)

    job_title = job_posting["job_title"]
    job_summary = job_posting.get("job_description", "")
    job_tech_skills = job_posting.get("technical_skills", [])
    job_soft_skills = job_posting.get("soft_skills", [])
    job_years_of_experience = job_posting.get("years_of_experience_required", "Not specified")
    job_requirements = job_posting.get("requirements", [])
    job_responsibilities = job_posting.get("responsibilities", [])

    resume_name = resume.get("personal_information", {}).get("name", "Unknown")
    resume_summary = resume.get("professional_summary", "")
    resume_years_experience = resume.get("years_of_experience", "")
    resume_skills = resume.get("technical_skills", [])
    resume_soft_skills = resume.get("soft_skills", [])
    resume_education = resume.get("education", [])
    resume_experience = resume.get("work_experience", [])

    system_instructions = """
    You are an expert Applicant Tracking System (ATS) evaluator.

    Your job is to **strictly analyze** how well a resume matches a job posting.

    You must identify the **technical skills**, **soft skills**, and **keywords** that are explicitly mentioned in the job posting and then determine which of those appear clearly in the resume.

    ---

    ## INSTRUCTIONS (Strict):

    1. The **universe of skills and keywords** must come **only** from the job posting.
    2. **Matching technical skills** = appear in both the job posting and the resume.
    3. **Missing technical skills** = appear in the job posting but **not** in the resume.
    4. **Matching keywords** = appear in both the job posting and the resume.
    5. **Missing keywords** = appear in the job posting but **not** in the resume.
    6. **Do NOT infer** technologies, skills, or synonyms not explicitly mentioned in the job posting.
    7. A skill only counts as matched if it appears clearly in the resume (in the skills section, summary, or experience).
    8. Be precise with terms:
    - “GitHub Actions” may match “GitHub CI/CD” if clearly implied.
    - “AWS Lambda” ≠ “Lambda”. AWS services must be explicitly named.
    - Distinguish between “GitHub” (version control) and “GitHub Actions” (CI/CD).
    9. Even if the ATS score is high, always include a **“recommendations”** section with:
    - Suggestions for improvement.
    - Gaps or partial matches.
    - Phrasing tips for clearer alignment.

    ---

    ## MATCHING RULES:

    - Convert all terms to lowercase.
    - Remove special characters (parentheses, commas, etc.).
    - Do not group terms (e.g., use “AWS Lambda”, “AWS SNS”, not “AWS (Lambda, SNS)”).
    - Do not duplicate items.
    - Match only if the context in the resume supports it clearly.

    ---

    ## EVALUATION CRITERIA:

    1. Technical Skills Match  
    2. Soft Skills Match  
    3. Experience Relevance & Duration  
    4. Responsibility Alignment  
    5. Professional Summary Fit  
    6. Keyword Match  
    7. Education Match  
    8. Cultural/Organizational Fit

    ---

    ## OUTPUT FORMAT (JSON ONLY):

    ```json
    {
    "ats_score": <integer from 0 to 100>,
    "matching_technical_skills": [list],
    "missing_technical_skills": [list],
    "matching_soft_skills": [list],
    "missing_soft_skills": [list],
    "keywords_matched": [list],
    "keywords_missing": [list],
    "years_of_experience_match": "Yes" or "No",
    "education_match": "Yes" or "No",
    "summary_match": "Strong" | "Partial" | "No",
    "responsibility_alignment": "Strong" | "Partial" | "No",
    "recommendations": [list]
    }

        """

    prompt = f"""
    ======= JOB POSTING =======
    Job Title: {job_title}
    Summary: {job_summary}
    Required Technical Skills: {job_tech_skills}
    Required Soft Skills: {job_soft_skills}
    Minimum Years of Experience: {job_years_of_experience}
    Required Education: {job_requirements}
    Key Responsibilities:
    {job_responsibilities}

    ======= RESUME =======
    Name: {resume_name}
    Summary: {resume_summary}
    Years of Experience: {resume_years_experience}
    Technical Skills: {resume_skills}
    Soft Skills: {resume_soft_skills}
    Education: {resume_education}
    Work Experience:
    {resume_experience}
    """

    genai.configure(api_key=your_api_key)
    model = genai.GenerativeModel(
        model_gemini,
        system_instruction=system_instructions,
        generation_config={
        "temperature": 0.0,
        "top_p": 1.0,
        "top_k": 1,
        "max_output_tokens": 1024
    }
    )

    response = model.generate_content(prompt)

    response_text = response.text.strip()

    if not response_text:
        print("❌ Empty response from Gemini.")
        return

    response_clean = (
        response_text.strip("`").replace("json", "").replace("JSON", "").strip()
    )

    try:
        result_json = json.loads(response_clean)
    except json.JSONDecodeError as e:
        print("❌ Error decoding JSON:", e)
        print("🔍 Raw Gemini output:")
        print(response_clean)
        return

    output_filepath = f"resume/ats_score_evaluation_pre.json"
    with open(output_filepath, "w", encoding="utf-8") as file:
        json.dump(result_json, file, ensure_ascii=False, indent=4)
        print(f"✅ Output saved to '{output_filepath}'")

def ats_score_evaluation_post():

    # Load files
    with open("resume/resume_final_to_word.json", "r", encoding="utf-8") as file:
        resume = json.load(file)

    with open("resume/ats_score_evaluation_pre.json", "r", encoding="utf-8") as file:
        ats_pre = json.load(file)

    with open("resume/job_posting.json", "r", encoding="utf-8") as file:
        job_posting = json.load(file)

    # Extract fields
    job_title = job_posting.get("job_title", "")
    job_summary = job_posting.get("job_description", "")
    job_tech_skills = job_posting.get("technical_skills", [])
    job_soft_skills = job_posting.get("soft_skills", [])
    job_years_of_experience = job_posting.get("years_of_experience_required", "Not specified")
    job_requirements = job_posting.get("requirements", [])
    job_responsibilities = job_posting.get("responsibilities", [])

    resume_name = resume.get("personal_information", {}).get("name", "Unknown")
    resume_summary = resume.get("professional_summary", "")
    resume_years_experience = resume.get("years_of_experience", "")
    resume_skills = resume.get("skills", [])
    resume_education = resume.get("education", [])
    resume_experience = resume.get("work_experience", [])

    # Previous match context
    previous_match_context = f"""
    ======= PREVIOUS MATCHES CONTEXT =======
    The following technical and soft skills were already validated as matched in a previous ATS evaluation. 
    Do not mark them as missing unless you can clearly prove they are not present in the resume or are no longer required in the new job posting.

    Previously matched technical skills: {ats_pre.get("matching_technical_skills", [])}
    Previously matched soft skills: {ats_pre.get("matching_soft_skills", [])}
    Previously matched keywords: {ats_pre.get("keywords_matched", [])}
    ========================================
    """

    previous_missing_context = f"""
    ======= CONTEXT: ITEMS PREVIOUSLY MISSING =======
    Please re-evaluate the following missing items. 
    IMPORTANT: Do NOT mark a skill or keyword as “matched” unless it is clearly and explicitly present in the resume text (summary, skills, or experience). Do NOT infer based on related terms or assumptions. Re-evaluate previously missing items strictly.

    Previously missing technical skills: {ats_pre.get("missing_technical_skills", [])}
    Previously missing soft skills: {ats_pre.get("missing_soft_skills", [])}
    Previously missing keywords: {ats_pre.get("keywords_missing", [])}
    ========================================
    """

    # System instructions
    system_instructions = """You are an expert Applicant Tracking System (ATS) evaluator.
        Your job is to strictly analyze how well a resume matches a job posting.

        You must infer skills and responsibilities only if they are clearly described in the resume. Be extremely accurate and consistent. This process is used to prepare a professional report and document, so clarity and formatting are critical.

        ## VERY IMPORTANT INSTRUCTIONS:

        1. Match ONLY the skills and keywords that appear in the job posting.
        2. Do NOT infer skills, technologies, or knowledge from the resume unless the job posting explicitly requires them.
        3. Do NOT include related or synonymous skills that are not directly listed in the job posting.
        4. A skill can only be considered "matched" if:
        - It is explicitly present in the job posting, and
        - It is clearly demonstrated in the resume (via skills list, professional summary, or experience).
        - "GitHub Actions" may match "GitHub CI/CD" if clearly implied.
        - AWS-specific services must be explicitly named (e.g., "AWS Lambda" ≠ just "Lambda").
        - Clearly distinguish between general terms like "GitHub" (version control) and specific tools like "GitHub Actions" (CI/CD).
        5. Even if the ATS score is high, **always include a “Recommendations” section**. These should:
        - Highlight areas that could be improved or better emphasized.
        - Indicate any gaps or partial matches in experience or responsibilities.
        - Suggest improvements in phrasing or contextualization (e.g., how to frame "client-facing" experience).
        6. If something is only partially covered (e.g., “responsibility alignment: partial”), explain why.


        ## Matching Rules (Strict and Normalized):

        - Normalize all terms by converting them to lowercase and trimming whitespace.
        - Remove special characters such as parentheses, commas, or symbols.
        - Do NOT use grouped terms like "aws (lambda, sns)" — expand each one as its own skill.
        - DO NOT include the same item in both the "matching" and "missing" lists.
        - Deduplicate all lists.
        - Match technical and soft skills only if context clearly supports them.

        ## EVALUATION CRITERIA (Very Strict):

        1. Technical Skills Match
        2. Soft Skills Match
        3. Experience Relevance & Duration
        4. Responsibility Alignment
        5. Professional Summary Fit
        6. Keyword Match
        7. Education Match
        8. Cultural/Organizational Fit

        ## OUTPUT FORMAT (JSON only):
        {
            "ats_score": <integer from 0 to 100>,
            "matching_technical_skills": [list],
            "missing_technical_skills": [list],
            "matching_soft_skills": [list],
            "missing_soft_skills": [list],
            "keywords_matched": [list],
            "keywords_missing": [list],
            "years_of_experience_match": "Yes" or "No",
            "education_match": "Yes" or "No",
            "summary_match": "Strong", "Partial", or "No",
            "responsibility_alignment": "Strong", "Partial", or "No",
            "recommendations": [list]
        }"""

    # Prompt to model
    prompt = f"""
    {previous_match_context}
    {previous_missing_context}
    ======= JOB POSTING =======
    Job Title: {job_title}
    Summary: {job_summary}
    Required Technical Skills: {job_tech_skills}
    Required Soft Skills: {job_soft_skills}
    Minimum Years of Experience: {job_years_of_experience}
    Required Education: {job_requirements}
    Key Responsibilities:
    {job_responsibilities}

    ======= RESUME =======
    Name: {resume_name}
    Summary: {resume_summary}
    Years of Experience: {resume_years_experience}
    Technical Skills: {resume_skills}
    Soft Skills: {resume_skills}
    Education: {resume_education}
    Work Experience:
    {resume_experience}
    """

    genai.configure(api_key=your_api_key)
    model = genai.GenerativeModel(
        model_gemini,
        system_instruction=system_instructions,
        generation_config={
        "temperature": 0.0,
        "top_p": 1.0,
        "top_k": 1,
        "max_output_tokens": 1024
    }
    )

    response = model.generate_content(prompt)

    response_text = response.text.strip()

    if not response_text:
        print("❌ Empty response from Gemini.")
        return

    response_clean = (
        response_text.strip("`").replace("json", "").replace("JSON", "").strip()
    )

    try:
        result_json = json.loads(response_clean)
    except json.JSONDecodeError as e:
        print("❌ Error decoding JSON:", e)
        print("🔍 Raw Gemini output:")
        print(response_clean)
        return

    output_filepath = f"resume/ats_score_evaluation_post.json"
    with open(output_filepath, "w", encoding="utf-8") as file:
        json.dump(result_json, file, ensure_ascii=False, indent=4)
        print(f"✅ Output saved to '{output_filepath}'")


def clean_text(text):
    if not isinstance(text, str) or not text.strip():
        return ""

    text = text.lower()  # Convert to lowercase
    text = re.sub(r'\d+', '', text)  # Remove numbers
    text = text.translate(str.maketrans("", "", string.punctuation))  # Remove punctuation
    text = re.sub(r'\s+', ' ', text).strip()  # Remove extra spaces

    # Remove stopwords
    stop_words = set(stopwords.words("english"))
    words = text.split()
    text = " ".join([word for word in words if word not in stop_words])

    return text



def clean_text_with_lemmatization(text):
    text = clean_text(text)
    words = text.split()
    return " ".join([lemmatizer.lemmatize(word) for word in words])

def normalize_keywords(raw_keywords_set):
    normalized_keywords = set()

    for kw in raw_keywords_set:
        words = kw.strip().lower().split()

        if len(words) == 1:
            normalized_keywords.add(words[0])
        elif len(words) >= 2:
            sorted_bigram = " ".join(sorted(words[:2]))
            normalized_keywords.add(sorted_bigram)

    return normalized_keywords


def jaccard_similarity(set1, set2):
    if not set1 or not set2:
        return 0.0
    intersection = len(set1.intersection(set2))
    union = len(set1.union(set2))
    return intersection / union

def join_all_resume_json():
    input_filepath = "resume/resume_education_info_personal.json"
    with open(input_filepath, "r", encoding="utf-8") as file_load:
       education_personal = json.load(file_load)

    input_filepath = "resume/resume_summary.json"
    with open(input_filepath, "r", encoding="utf-8") as file_load:
       summary = json.load(file_load)

    input_filepath = var_resume_match_skils
    with open(input_filepath, "r", encoding="utf-8") as file_load:
       skills_json = json.load(file_load)

    # Create a set to avoid duplicates
    combined_skills = set(skills_json["technical_skills"] + skills_json["soft_skills"])

    input_filepath = "resume/resume_user_answers.json"
    if os.path.exists(input_filepath):
        with open(input_filepath, "r", encoding="utf-8") as file_load:
            user_answers = json.load(file_load)

        # Add user_answers skills
        for entry in user_answers:
            print(entry)
            skill = entry.get("skill")
            if isinstance(skill, str):
                combined_skills.add(skill)


    # Convert the set to a list and structure it in the new JSON
    final_skills_json = {"skills": list(combined_skills)}

    input_filepath = "resume/resume_final_experience.json"
    with open(input_filepath, "r", encoding="utf-8") as file_load:
       final_experience = json.load(file_load)

    final_resume_json = {
        "education": education_personal.get("education", []),
        "personal_information": education_personal.get("personal_information", []),
        "professional_summary": summary.get("professional_summary", []),
        "skills": final_skills_json.get("skills", []),
        "work_experience": final_experience.get("work_experience",[])
    }

    # Save the updated JSON file
    output_file = "resume/resume_final_to_word.json"
    with open(output_file, "w", encoding="utf-8") as file:
        json.dump(final_resume_json , file, indent=4, ensure_ascii=False)

def validate_with_gemini(skill, detail):
    validation_prompt = (
        f"Evaluate the following response regarding experience with {skill}.\n"
        f"Ensure it includes:\n"
        f"- A clear explanation of how the experience was obtained.\n"
        f"- At least one action verb describing what was done.\n"
        f"- At least one quantifiable metric or measurable impact. This can include:\n"
        f"  - A percentage improvement (e.g., 'Increased efficiency by 25%').\n"
        f"  - A numerical value (e.g., 'Led a team of 10 engineers').\n"
        f"  - A frequency or time-based metric (e.g., 'Completed weekly reports over a 6-month period').\n"
        f"  - A quantity or scope (e.g., 'Managed inventory of 1,200 items').\n\n"
        f"Response to evaluate:\n{detail}\n\n"
        f"### Evaluation Criteria ###\n"
        f"- If the response meets all criteria, return: 'Evaluation: ✅ Strong response.'\n"
        f"- If the response is missing details, return: 'Evaluation: ❌ Needs improvement.' "
        f"and provide a rewritten version of the response following a strong format.\n\n"
        f"### Examples of strong responses ###\n"
        f"'Implemented a predictive maintenance system using Python, reducing machine downtime by 25% over six months.'\n"
        f"'Led a team of 10 engineers in developing an AI-driven analytics tool.'\n"
        f"'Led the team through a seven-day sales event, successfully aligning production schedules with demand.'\n"
        f"'Controlled a schedule of approximately 15 weekly sales visits to assigned clients, enhancing relationships and communication skills.'\n"
        f"'Managed inventory of over 1,200 items across 3 warehouses, improving supply chain visibility.'\n"
        f"'Delivered training sessions to more than 40 employees over a 3-month onboarding program.'\n\n"
        f"Now, if the response needs improvement, provide a corrected version formatted as:\n"
        f"'Example: [Your improved response here]'"
    )
    
    genai.configure(api_key = your_api_key)
    model = genai.GenerativeModel(
    model_gemini,
    system_instruction=validation_prompt,
        generation_config={
        "temperature": 0.0,
        "top_p": 1.0,
        "top_k": 1,
        "max_output_tokens": 1024
    }
    )

    try:
        response = model.generate_content(validation_prompt)
        feedback = response.text.strip()

        if "✅ Strong response" in feedback:
            return True, feedback
        elif "❌ Needs improvement" in feedback:
            improved_response_match = re.search(r"Example:\s*(.+)", feedback, re.IGNORECASE)
            if improved_response_match:
                example = improved_response_match.group(1).strip()
            else:
                example = "Ensure you include a percentage or numerical value, such as 'Reduced processing time by 30%' or 'Led a team of 5 engineers'."
            return False, example
    except Exception as e:
        print(f"Error communicating with Gemini: {e}")
        return False, "Ensure you include a percentage or numerical value."


def resume_skills():
    input_filepath = var_resume_json

    with open(input_filepath, "r", encoding="utf-8") as file_load:
       cv_data = json.load(file_load)

    input_filepath = f"resume/job_posting.json"

    with open(input_filepath, "r", encoding="utf-8") as file_load:
       job_data = json.load(file_load)

    cv_technical_skills = set(cv_data.get("technical_skills", []))
    job_technical_skills = set(job_data.get("technical_skills", []))

    cv_soft_skills = set(cv_data.get("soft_skills", []))
    job_soft_skills = set(job_data.get("soft_skills", []))

    missing_technical_skills = list(job_technical_skills - cv_technical_skills)[:5]
    missing_soft_skills = list(job_soft_skills - cv_soft_skills)[:5]

    match_technical_skills = list(cv_technical_skills & job_technical_skills)
    match_soft_skills = list(cv_soft_skills & job_soft_skills)

    missing_skills = {
        "technical_skills": missing_technical_skills,
        "soft_skills": missing_soft_skills
    }

    match_skills = {
        "technical_skills": match_technical_skills,
        "soft_skills": match_soft_skills
    }

    with open("resume/resume_missing_skills.json", "w") as file:
        json.dump(missing_skills, file, indent=4)
    print("Skills missing saved in 'resume/resume_missing_skills.json'.")

    with open(var_resume_match_skils, "w") as file:
        json.dump(match_skills, file, indent=4)
    print("Skills missing saved in 'resume/resume_match_skills.json'.")



def resume_education_info_personal():

    input_filepath = var_resume_json
    with open(input_filepath, "r", encoding="utf-8") as file_load:
       original_cv = json.load(file_load)
    
    output_file = {
        "personal_information": original_cv.get("personal_information", {}),
        "education": original_cv.get("education", {})
    }


    output_filepath = f"resume/resume_education_info_personal.json"
    with open(output_filepath, "w", encoding="utf-8") as file_save:
        json.dump(output_file, file_save, ensure_ascii=False, indent=4)
        print(f"Output saved to '{output_filepath}'.")

def resume_promt_summary():

    input_filepath = var_resume_json
    with open(input_filepath, "r", encoding="utf-8") as file_load:
       resume = json.load(file_load)
    old_summary = resume["professional_summary"]
    education = resume["education"]
    year_experience = resume["years_of_experience"]
    
    input_filepath = f"resume/resume_updated.json"
    with open(input_filepath, "r", encoding="utf-8") as file_load:
       resume = json.load(file_load)
    experience_updated = resume

    # Create final skills json 
    input_filepath = f"resume/job_posting.json"
    with open(input_filepath, "r", encoding="utf-8") as file_load:
       job_offer = json.load(file_load)
    
    input_filepath = var_resume_match_skils
    with open(input_filepath, "r", encoding="utf-8") as file_load:
       skills_json = json.load(file_load)

    # Create a set to avoid duplicates
    combined_skills = set(skills_json["technical_skills"] + skills_json["soft_skills"])

    input_filepath = "resume/resume_user_answers.json"
    if os.path.exists(input_filepath):
        with open(input_filepath, "r", encoding="utf-8") as file_load:
            user_answers = json.load(file_load)

        # Add user_answers skills
        for entry in user_answers:
            print(entry)
            skill = entry.get("skill")
            if isinstance(skill, str):
                combined_skills.add(skill)


    # Convert the set to a list and structure it in the new JSON
    final_skills_json = {"skills": list(combined_skills)}

    system_instructions ="""
    You are an HR specialist skilled in processing and analyzing resumes.
    Your task is to generate a concise and professional summary strictly based on the provided resume.

    **Instructions:**
    - Use only the information available in the resume. Do not infer or add any details that are not explicitly mentioned.
    - Focus on highlighting relevant experience, skills, and qualifications.
    - Ensure clarity, coherence, and alignment with the job offer.
    - The response must be in JSON format only, without any explanations or additional text.
    - Rewrite the professional summary to align with the job posting’s focus on contributing to projects that support long-term sustainability and global investment strategies. Use clear, action-oriented language, and highlight relevant skills or experience that demonstrate the candidate’s ability to contribute to sustainable initiatives or global impact.

    **Output Format:**
    {
        "professional_summary": "Generated summary here"
    }
    """

    genai.configure(api_key = your_api_key)
    model = genai.GenerativeModel(
    model_gemini,
    system_instruction=system_instructions,
    )

    response = model.generate_content(f"The old resume to analyze is {old_summary} the education is {education}, the years of experience is {year_experience}, the new experience updated is {experience_updated}, and the updated skills are {final_skills_json},  and the job offer is {job_offer}")
    cleaned_response = response.text.strip(clean_json).strip("```").replace("\n", "")

    json_file = json.loads(cleaned_response)

    output_filepath = f"resume/resume_summary.json"
    with open(output_filepath, "w", encoding="utf-8") as file_save:
        json.dump(json_file, file_save, ensure_ascii=False, indent=4)
        print(f"Output saved to '{output_filepath}'.")


def resume_delete_experience_not_related():

    input_filepath = var_resume_json
    with open(input_filepath, "r", encoding="utf-8") as file_load:
       resume = json.load(file_load)
    
    job_experience = {"work_experience":resume.get("work_experience", {})}
    
    input_filepath = f"resume/job_posting.json"
    with open(input_filepath, "r", encoding="utf-8") as file_load:
       job_offer = json.load(file_load)
    
    system_instructions = """
    You are an HR specialist skilled in processing and analyzing resumes. Your task is to analyze each achievement in the work experience section and remove those that are not relevant to the given job offer.

    **Instructions:**
    - Use only the information available in the resume. Do not infer or add any details that are not explicitly mentioned.
    - Evaluate each achievement individually and determine if the experience and skills it describes align with the job posting.
    - Remove achievements that are not relevant to the job offer.
    - Maintain the original JSON format, ensuring that only the non-relevant achievements are removed.
    - Do not provide any explanations or extra text, only return the modified JSON.

    **Output Format:**
    (Same as input, but with non-relevant achievements removed)
    """

    genai.configure(api_key = your_api_key)
    model = genai.GenerativeModel(
    model_gemini,
    system_instruction=system_instructions,
    )

    response = model.generate_content(f"The work experience seccion to analyze is {job_experience} and the job offer is {job_offer}")
    cleaned_response = response.text.strip(clean_json).strip("```").replace("\n", "")
    json_file = json.loads(cleaned_response)
    # Save the result to the output file
    output_filepath = f"resume/resume_delete_experience_not_relate.json"
    with open(output_filepath, "w", encoding="utf-8") as file_save:
        json.dump(json_file, file_save, ensure_ascii=False, indent=4)
        print(f"Output saved to '{output_filepath}'.")


def customize_cv() -> dict:
    """
    Customizes the CV based on the extracted CV data and the job offer data.
    Both inputs are dictionaries parsed from JSON.
    
    The output JSON should include all important personal details (Name, Email, Phone, SocialMedia)
    along with a structured breakdown:
      - "Summary": <Rewritten summary text>,
      - "Skills": <Rewritten skills text>,
      - "Experience": [ { "Company": <Company name>, "Dates": <Start date - End date>, "Functions": <Rewritten description> }, ... ],
      - "Education": [ { "Institution": <Institution name>, "Dates": <Start date - End date>, "Degree": <Degree obtained> }, ... ]
      
    Do not invent new content; only enhance and reorganize the existing information.
    """


    input_filepath = var_resume_json
    with open(input_filepath, "r", encoding="utf-8") as file_load:
        original_cv = json.load(file_load)


    input_filepath = f"resume/job_posting.json"
    with open(input_filepath, "r", encoding="utf-8") as file_load:
        job_description = json.load(file_load)\
    
    prompt = f"""
    Using the following inputs, generate a revised version of the CV that improves the wording and structure to clearly emphasize the candidate's relevant skills and experiences for the job offer.
    Do not invent any new content; only enhance and reorganize the existing information.

    Original CV:
    {original_cv}

    Job Offer:
    {job_description}

    Please output the rewritten CV in a structured JSON format with the following keys:

    {{
    "PersonalInfo": {{
        "Name": "<Candidate's full name>",
        "Address": "<Candidate's location Address>",
        "Email": "<Candidate's email address>",
        "Phone": "<Candidate's phone number>",
        "SocialMedia": "<Links to social media profiles or other contact details>"
    }},
    "Summary": "<Rewritten summary text>",
    "Skills": "<Rewritten skills text>",
    "Experience": [
        {{
        "Company": "<Company name>",
        "Dates": "<Start date - End date>",
        "Functions": "<Rewritten description of responsibilities and functions>"
        }}
    ],
    "Education": [
        {{
        "Institution": "<Institution name>",
        "Dates": "<End date>",
        "Degree": "<Degree obtained>"
        }}
    ]
    }}

    Ensure the JSON is valid.
    """
    genai.configure(api_key = your_api_key)
    model = genai.GenerativeModel(
    model_gemini
    )
    response = model.generate_content(prompt)

    cleaned_response = response.text.strip(clean_json).strip("```").replace("\n", "")

    json_file = json.loads(cleaned_response)

    output_filepath = f"resume/resume_customization.json"
    with open(output_filepath, "w", encoding="utf-8") as file_save:
        json.dump(json_file, file_save, ensure_ascii=False, indent=4)
        print(f"Output saved to '{output_filepath}'.")

def extract_cv_information(uploaded_pdf):
    # Read the PDF content using pypdf
    pdf_reader = PdfReader(BytesIO(uploaded_pdf.getvalue()))
    pdf_text = ""
    for page in pdf_reader.pages:
        pdf_text += page.extract_text()

    system_instructions = """
    You are an HR analyst responsible for processing and analyzing resumes. Your task is to extract the following information from a given resume:

    1. Technical Skills:
        -Identify and list all technical skills explicitly mentioned in the resume.
        -These may include programming languages, frameworks, libraries, databases, cloud platforms, and relevant tools.
    2. Soft Skills:
        -Identify and list all soft skills mentioned in the resume.
        -Additionally, infer soft skills based on the candidate's experience, achievements, and responsibilities.
        -Examples include leadership, teamwork, adaptability, problem-solving, communication, critical thinking, and emotional intelligence.
    3. Years of Experience:
        -Calculate and provide the total number of years of professional experience.
    4. Education:
        -Extract and list all education details, including:
            -Degree
            -Institution
            -Location
            -Start date
            -End date
    5. Experience Level:
        -Determine the candidate's highest level of experience based on job roles, industries, and responsibilities.
    6. Work Experience:
        -Extract all previous job positions and structure them as a list of dictionaries, where each dictionary represents a job.
        -Each dictionary must contain:
            -Job title
            -Company
            -Location
            -Start date
            -End date
            -Achievements and responsibilities associated with that role
    7. Professional Summary:
        -Extract the professional summary from the resume.
    8. Languages & Proficiency (Optional: If not found, omit this section):
        -Identify all languages spoken by the candidate and their proficiency level.
        -If no language information is found, exclude this section from the final JSON.
    9. Certifications (Optional: If not found, omit this section):
        -List all certifications the candidate has obtained, including:
        -Certification name
        -Issuing organization
        -Year issued
        -If no certifications are found, exclude this section from the final JSON.
    10. Personal information:
        -Name
        -Phone
        -Email
        -Addres
        -Identify and list Social media links

    11. Output Format:
        -Return only the JSON response in the exact structure below.
        -Do not include any explanations, extra text, or formatting beyond the JSON itself.

        {
        "technical_skills": [A],
        "soft_skills": [B],
        "years_of_experience": C,
        "education": [
            {
            "degree": D,
            "institution": D,
            "location": D,
            "start_date":D, 
            "end_date":D 
            }
        ],
        "experience_level": E,
        "work_experience": [
            {
            "job_title":F ,
            "company": F,
            "location": F,
            "start_date":F,
            "end_date":F ,
            "key": "company-job_title"
            "achievement": [F
            ]
            }
        ],
        "professional_summary": G,
        "languages": [
            {
            "language": H,
            "proficiency": H
            }
        ],
        "certifications": [
            {
            "name": I,
            "issuing_organization":I,
            "year_issued": I
            }
        ],
        "personal_information":{
            "name": J,
            "phone": J,
            "email": J,
            "addres": J,
            "social_media":[J]
        }
        }
    """

    genai.configure(api_key = your_api_key)
    model = genai.GenerativeModel(
    model_gemini,
    system_instruction=system_instructions,
    )

    response = model.generate_content(f"The resume to analyze is {pdf_text}")
    cleaned_response = response.text.strip(clean_json).strip("```").replace("\n", "")

    json_file = json.loads(cleaned_response)
    # Save the result to the output file
    output_filepath = var_resume_json
    with open(output_filepath, "w", encoding="utf-8") as file_save:
        json.dump(json_file, file_save, ensure_ascii=False, indent=4)
        print(f"Output saved to '{output_filepath}'.")

def extract_job_posting_information(uploaded_job):
    # Read the PDF content using pypdf
    pdf_reader = PdfReader(BytesIO(uploaded_job.getvalue()))
    pdf_text = ""
    for page in pdf_reader.pages:
        pdf_text += page.extract_text()

    system_instructions = """
    You are a job market analyst specializing in processing and analyzing job postings. 
    Your task is to extract key information from a given job listing and structure it into a JSON format. 
    The extracted data should include:

    1. Job Information:
        -The job title of the position.
    2. Company Information:
        -The company name offering the position.
        -The industry in which the company operates.
        -The location where the job is based.
        -Whether the job is remote-friendly (true/false).
    3. Employment Type & Experience Level:
        -The employment type (e.g., Full-time, Part-time, Contract, Internship).
        -The experience level required (e.g., Junior, Mid, Senior, Lead).
        -The minimum years of experience required.
    4. Salary Information:
        -The salary range (minimum and maximum salary).
        -The currency in which the salary is offered.
        -The payment frequency (e.g., Hourly, Monthly, Annual).
    5. Job Description:
        -A short summary of the job description provided in the job posting.
    6. Responsibilities:
        -A list of job responsibilities as mentioned in the posting.
    7. Requirements:
        -A list of all mandatory job requirements specified in the posting.
    8. Skills Required:
        -Identify and list all technical skills required for the job.
        -Identify and list all soft skills required or implied in the posting.
    9. Education & Certifications (Optional: If not found, omit these sections):
        -Identify the minimum education level required (degree and field of study).
        -List any preferred institutions mentioned in the job description.
        -Identify any required certifications, including the certification name and issuing organization.
    10. Language Requirements (Optional: If not found, omit this section):
        -Identify all languages required for the job, along with the proficiency level.
    11. Benefits & Perks (Optional: If not found, omit this section):
        -Extract and list all benefits mentioned in the job posting (e.g., health insurance, remote work, bonuses, flexible hours).
    12. Application Information:
        -The application deadline (if provided).
        -The job posting date (if available).
        -The official application link where candidates can apply.

    13. Output Format:
        -Return ONLY the JSON response in the exact structure below.
        -Do NOT include any explanations, comments, or extra text beyond the JSON itself.

        {
        "job_title": "A",
        "company": {
            "name": "B",
            "industry": "C",
            "location": "D",
            "remote": E
        },
        "employment_type": "F",
        "experience_level": "G",
        "years_of_experience_required": H,
        "salary_range": {
            "min": I,
            "max": J,
            "currency": "K",
            "payment_frequency": "L"
        },
        "job_description": "M",
        "responsibilities": [
            "N"
        ],
        "requirements": [
            "O"
        ],
        "technical_skills": [
            "P"
        ],
        "soft_skills": [
            "Q"
        ],
        "education_required": [
            {
            "degree": "R",
            "field_of_study": "S",
            "preferred_institutions": ["T"]
            }
        ],
        "certifications_required": [
            {
            "name": "U",
            "issuing_organization": "V"
            }
        ],
        "languages_required": [
            {
            "language": "W",
            "proficiency": "X"
            }
        ],
        "benefits": [
            "Y"
        ],
        "application_deadline": "Z",
        "job_posting_date": "AA",
        "application_link": "BB"
        }
    """

    genai.configure(api_key = your_api_key)
    model = genai.GenerativeModel(
    model_gemini,
    system_instruction=system_instructions,
    )

    response = model.generate_content(f"The job posting to analyze is {pdf_text}")
    cleaned_response = response.text.strip(clean_json).strip("```").replace("\n", "")

    json_file = json.loads(cleaned_response)
    output_filepath = f"resume/job_posting.json"
    with open(output_filepath, "w", encoding="utf-8") as file_save:
        json.dump(json_file, file_save, ensure_ascii=False, indent=4)
        print(f"Output saved to '{output_filepath}'.")



def extract_job_posting_information_from_str(uploaded_job):

    pdf_text = uploaded_job

    system_instructions = """
    You are a job market analyst specializing in processing and analyzing job postings. 
    Your task is to extract key information from a given job listing and structure it into a JSON format. 
    The extracted data should include:

    1. Job Information:
        -The job title of the position.
    2. Company Information:
        -The company name offering the position.
        -The industry in which the company operates.
        -The location where the job is based.
        -Whether the job is remote-friendly (true/false).
    3. Employment Type & Experience Level:
        -The employment type (e.g., Full-time, Part-time, Contract, Internship).
        -The experience level required (e.g., Junior, Mid, Senior, Lead).
        -The minimum years of experience required.
    4. Salary Information:
        -The salary range (minimum and maximum salary).
        -The currency in which the salary is offered.
        -The payment frequency (e.g., Hourly, Monthly, Annual).
    5. Job Description:
        -A short summary of the job description provided in the job posting.
    6. Responsibilities:
        -A list of job responsibilities as mentioned in the posting.
    7. Requirements:
        -A list of all mandatory job requirements specified in the posting.
    8. Skills Required:
        -Identify and list all technical skills required for the job.
        -Identify and list all soft skills required or implied in the posting.
    9. Education & Certifications (Optional: If not found, omit these sections):
        -Identify the minimum education level required (degree and field of study).
        -List any preferred institutions mentioned in the job description.
        -Identify any required certifications, including the certification name and issuing organization.
    10. Language Requirements (Optional: If not found, omit this section):
        -Identify all languages required for the job, along with the proficiency level.
    11. Benefits & Perks (Optional: If not found, omit this section):
        -Extract and list all benefits mentioned in the job posting (e.g., health insurance, remote work, bonuses, flexible hours).
    12. Application Information:
        -The application deadline (if provided).
        -The job posting date (if available).
        -The official application link where candidates can apply.

    13. Output Format:
        -Return ONLY the JSON response in the exact structure below.
        -Do NOT include any explanations, comments, or extra text beyond the JSON itself.

        {
        "job_title": "A",
        "company": {
            "name": "B",
            "industry": "C",
            "location": "D",
            "remote": E
        },
        "employment_type": "F",
        "experience_level": "G",
        "years_of_experience_required": H,
        "salary_range": {
            "min": I,
            "max": J,
            "currency": "K",
            "payment_frequency": "L"
        },
        "job_description": "M",
        "responsibilities": [
            "N"
        ],
        "requirements": [
            "O"
        ],
        "technical_skills": [
            "P"
        ],
        "soft_skills": [
            "Q"
        ],
        "education_required": [
            {
            "degree": "R",
            "field_of_study": "S",
            "preferred_institutions": ["T"]
            }
        ],
        "certifications_required": [
            {
            "name": "U",
            "issuing_organization": "V"
            }
        ],
        "languages_required": [
            {
            "language": "W",
            "proficiency": "X"
            }
        ],
        "benefits": [
            "Y"
        ],
        "application_deadline": "Z",
        "job_posting_date": "AA",
        "application_link": "BB"
        }
    """

    genai.configure(api_key = your_api_key)
    model = genai.GenerativeModel(
    model_gemini,
    system_instruction=system_instructions,
    )

    response = model.generate_content(f"The job posting to analyze is {pdf_text}")
    cleaned_response = response.text.strip(clean_json).strip("```").replace("\n", "")

    json_file = json.loads(cleaned_response)

    output_filepath = f"resume/job_posting.json"
    with open(output_filepath, "w", encoding="utf-8") as file_save:
        json.dump(json_file, file_save, ensure_ascii=False, indent=4)
        print(f"Output saved to '{output_filepath}'.")



def find_best_jobs(cv_file, job_offers):
    """Finds the top 10 best matching jobs using cosine similarity."""
    cv_text = extract_text_from_file(cv_file)
    
    job_texts = [job['Description'] for job in job_offers]
    job_titles = [job['Title'] for job in job_offers]
    
    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform([cv_text] + job_texts)
    
    similarities = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:]).flatten()
    best_matches_idx = similarities.argsort()[-10:][::-1]
    
    best_jobs = [
        {"Title": job_titles[i], "Similarity": similarities[i]} 
        for i in best_matches_idx
    ]
    
    return pd.DataFrame(best_jobs)


def skills_missing():
    # Extract skills from resume and job posting
    cv_skills = set(cv_data.get("technical_skills", []))
    job_skills = set(job_data.get("technical_skills", []))

    # Identify missing skills (limit to 5)
    missing_skills = list(job_skills - cv_skills)[:5]

    for skill in missing_skills:
        print(f"Do you have experience with {skill}? (yes/no)")
        answer = input("Your answer: ").strip().lower()

        if answer == "yes":
            while True:
                # Ask the user to describe their experience with the skill
                detail = input(f"Describe your experience with {skill}, including how you obtained it and a metric or result achieved: ")

                # Validate the response with Gemini
                validation_prompt = (
                    f"Evaluate the following response regarding experience with {skill}. "
                    f"Ensure it includes:\n"
                    f"- A clear explanation of how the experience was obtained.\n"
                    f"- At least one action verb describing what was done.\n"
                    f"- A quantifiable metric or measurable impact.\n\n"
                    f"Response to evaluate:\n{detail}\n\n"
                    f"### Evaluation Criteria ###\n"
                    f"- If the response meets all criteria, return: 'Evaluation: ✅ Strong response.'\n"
                    f"- If the response is missing details, return: 'Evaluation: ❌ Needs improvement.' "
                    f"and provide a suggestion with an example of how the response should be structured.\n\n"
                    f"### Example of a strong response ###\n"
                    f"'Implemented a predictive maintenance system using Python, reducing machine downtime by 25% over six months.'\n\n"
                )

                try:
                    response = model.generate_content(validation_prompt)
                    feedback = response.text.strip()
                    print("\n🔹 Gemini Response:\n", feedback)  # Display Gemini's response for debugging

                    # If the response is strong, save it and exit the loop
                    if re.search(r"✅ Strong response", feedback, re.IGNORECASE):
                        save_answers[skill] = detail
                        break

                    # If the response needs improvement, show feedback and an example
                    if re.search(r"❌ Needs improvement", feedback, re.IGNORECASE):
                        print("\n⚠️ Your answer needs improvement.")

                        # Extract a strong response example from Gemini's feedback
                        example_match = re.search(r"Example of a strong response:\s*(.+)", feedback, re.IGNORECASE)
                        if example_match:
                            print(f"\n🔄 Example of how you should structure your answer:\n{example_match.group(1).strip()}")

                        print("\nPlease try again with more detail using action verbs and metrics.")

                except Exception as e:
                    print(f"\n❌ Error communicating with Gemini: {e}")
                    print("Skipping this skill...")
                    break

    # Save responses to a JSON file
    with open("user_answers.json", "w") as file:
        json.dump(save_answers, file, indent=4)

    print("\n✅ Answers saved successfully.")



# Reference
# (OpenAI, ChatGPT o1, first prompt, 2025): I have this template in Word, this Json, both have the same keys, guide me to make a code that reeplace the information in the template with the info in Json
# (Claude, 3.5 Sonnet, last prompt, 2025): The template is not filling completly good, help me to correct the format

import re
import json
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def split_into_sentences(text):
    """Splits the text into sentences using punctuation."""
    print(text)
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    print(sentences)
    output = [s.strip() for s in sentences if s.strip()]
    print("split_into_sentences line", 844)
    return output
    

class CVGenerator:
    CUSTOM_BULLET_STYLE = 'Custom Bullet'
    
    def __init__(self, template_path):
        self.doc = Document(template_path)
        self.setup_styles()
        
    def setup_styles(self):
        """Set up document styles."""
        styles = self.doc.styles
        
        # Normal style
        style = styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        
        # Custom bullet style
        if self.CUSTOM_BULLET_STYLE not in styles:
            bullet_style = styles.add_style(self.CUSTOM_BULLET_STYLE, WD_STYLE_TYPE.PARAGRAPH)
            bullet_style.base_style = styles['Normal']
            bullet_style.paragraph_format.left_indent = Inches(0.25)
            bullet_style.paragraph_format.first_line_indent = Inches(-0.25)
            
    def add_bullet_paragraph(self, text):
        print("add_bullet_paragraph line", 871)
        """Add a bullet point paragraph."""
        paragraph = self.doc.add_paragraph()
        paragraph.style = self.CUSTOM_BULLET_STYLE
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        run_bullet = paragraph.add_run('• ')
        run_bullet.bold = True
        paragraph.add_run(text.strip())
        return paragraph
        
    def add_section_title(self, title):
        """Add a section title with proper formatting."""
        # Add space before section
        spacing_para = self.doc.add_paragraph()
        spacing_para.paragraph_format.space_before = Pt(12)
        
        # Section title
        paragraph = self.doc.add_paragraph()
        paragraph.add_run(title).bold = True
        self.add_horizontal_line(paragraph)
        
    def add_horizontal_line(self, paragraph):
        """Add a horizontal line after a section title."""
        p = paragraph._p
        p_pr = p.get_or_add_pPr()
        bottom_border = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom_border.append(bottom)
        p_pr.append(bottom_border)
        
    def add_right_aligned_text(self, paragraph, text):
        """Add right aligned text to a paragraph using tab stops."""
        paragraph.paragraph_format.tab_stops.clear_all()
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(6), WD_TAB_ALIGNMENT.RIGHT
        )
        paragraph.add_run('\t' + text)
        
    def fill_cv(self, data):
        """Fill the CV with the provided data."""
        # Clear the document
        for paragraph in self.doc.paragraphs[:]:
            p = paragraph._element
            p.getparent().remove(p)
            
        # Name
        name_paragraph = self.doc.add_paragraph()
        name_text = data.get('personal_information', {}).get('name', '')
        name_run = name_paragraph.add_run(name_text)
        name_run.bold = True
        name_run.font.size = Pt(48)

        # Contact information (including Social Media)
        personal_info = data.get('personal_information', {})
        address = personal_info.get('addres', '')
        phone = personal_info.get('phone', '')
        email = personal_info.get('email', '')
        social_list = personal_info.get('social_media', [])
        social = ", ".join(social_list) if isinstance(social_list, list) else social_list

        # Create a contact string that only includes non-empty fields
        contact_parts = [address, phone, email, social]
        contact = " | ".join([str(part) for part in contact_parts if part])
        
        contact_paragraph = self.doc.add_paragraph()
        contact_paragraph.add_run(contact)
        contact_paragraph.paragraph_format.space_after = Pt(12)
        
        # Summary (justified)
        self.add_section_title("Summary")
        summary_text = data.get('professional_summary', '')
        summary_para = self.doc.add_paragraph()
        summary_para.add_run(summary_text)
        summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        summary_para.paragraph_format.space_after = Pt(12)
        
        # Experience
        self.add_section_title("Experience")
        for exp in data.get('work_experience', []):
            achievements = exp.get('achievement', [])
            if not achievements or (isinstance(achievements, list) and not any(item.strip() for item in achievements)):
                continue
                
            exp_para = self.doc.add_paragraph()
            exp_para.paragraph_format.space_before = Pt(12)
            job_title = exp.get('job_title', '')
            company = exp.get('company', '')
            location = exp.get('location', '')
            exp_text = f"{job_title} | {company} | {location}"
            exp_run = exp_para.add_run(exp_text)
            exp_run.bold = True
            start_date = exp.get('start_date', '')
            end_date = exp.get('end_date', '')
            dates_text = f"{start_date} - {end_date}"
            self.add_right_aligned_text(exp_para, dates_text)
            
            # Functions (bullet points, justified)
            functions_text = exp.get('achievement', '')
            functions = functions_text
            for function in functions:
                if function.strip():
                    bullet_para = self.add_bullet_paragraph(function)
                    bullet_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    bullet_para.paragraph_format.space_after = Pt(3)
                    
        # Education
        self.add_section_title("Education")
        for edu in data.get('education', []):
            edu_para = self.doc.add_paragraph()
            edu_para.paragraph_format.space_before = Pt(6)
            institution = edu.get('institution', '')
            edu_run = edu_para.add_run(f"{institution}")
            edu_run.bold = True
            dates = edu.get('end_date', '')
            self.add_right_aligned_text(edu_para, dates)
            
            degree_para = self.doc.add_paragraph()
            degree_para.add_run(edu.get('degree', ''))
            degree_para.paragraph_format.space_after = Pt(6)
            
        # Skills in columns
        self.add_section_title("Skills")
        skills_list = data.get('skills', [])  # Obtener la lista directamente
        skills = [skill.strip() for skill in skills_list if isinstance(skill, str)]

        num_columns = 3  # Change to 2 if preferred
        num_skills = len(skills)
        num_rows = (num_skills + num_columns - 1) // num_columns
        
        table = self.doc.add_table(rows=num_rows, cols=num_columns)
        table.autofit = True
        
        skill_index = 0
        for r in range(num_rows):
            row_cells = table.rows[r].cells
            for c in range(num_columns):
                if skill_index < num_skills:
                    paragraph = row_cells[c].paragraphs[0]
                    run = paragraph.add_run('• ' + skills[skill_index])
                    run.font.size = Pt(11)
                    skill_index += 1
                else:
                    row_cells[c].text = ""
    
    def save(self, output_path):
        """Save the document to the specified path."""
        self.doc.save(output_path)
            
def generate_cv():
    json_file_path = f"resume/resume_final_to_word.json"
    template_path = f"template/template1.docx"
    """Generate CV from JSON data"""
    try:
        with open(json_file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
        
        user_name = data.get('personal_information', {}).get('name', 'Unknown').strip()
        user_name = " ".join(user_name.title().split())
        output_path = f"output/{user_name}_customization.docx"  

       # Create and save the Word version
        generator = CVGenerator(template_path)
        generator.fill_cv(data)
        generator.save(output_path)
        print(f"Word CV generated successfully: {output_path}")
        return True
        
    except Exception as e:
        print(f"Error generating CV: {str(e)}")
        return False

    


""" First Prompt in chatgpt
I need that when a resume in PDF format is read, a .json file is generated that can be recreated as best as possible. I want them to be {{
      "technical_skills": [A],
      "soft_skills": [B],
      "years_of_experience": C,
      "education": [
       {{
       "name": "",
       "place": "D",}}
       ],
       "experience_level": "E",
       "key_achievements": [
       {{
        "achievement": "F",
        "context": "F"
        }}
        ],
        "domain_expertise": [G],
        "job_titles": [H],  
        "industries": [I],
        "professional_summary": "J"
        }}
"""

""" Last Prompt in chatgpt
You are an HR analyst and you receive resumes and must analyze them to answer:
A. Identify and list all ***hard skills*** (e.g., programming languages, frameworks, libraries, databases, cloud platforms, and tools) explicitly mentioned in the resume.
B. Identify and list all ***soft skills*** (e.g., leadership, teamwork, adaptability, problem solving, communication, critical thinking, emotional intelligence) found in the resume and that you can also infer from the resume.
C. How many years of total professional experience does the candidate have?
D. Identify and list all educational information, degree, institution, location, start date, end date
E. What is the candidate's highest level of experience based on job roles, industries, and responsibilities?
F. Identify and list all of the candidate's work experience candidate:
- A list of dictionaries where in each dictionary there is a job that the candidate had.
- The dictionary should have the job title, company, location, start date, end date and all the achievements or responsibilities that the CV has in that particular position.
- The same information for all jobs
G. Find the professional summary.
H. All languages ​​and skills, if you don't find them in the candidate ignore and delete this section
I. All certifications that the candidate has, if you don't find them delete this section
J. Find the professional summary.

Returns only the JSON response in the exact structure shown below, without explanations or additional text:

{
"technical_skills": [A],
"soft_skills": [B],
"years_of_experience": C,
"education": [
{
"degree": D,
"institution": D,
"location": D,
"start_date":D,
"end_date":D
}
],
"experience_level": E,
"work_experience": [
{
"job_title":F ,
"company": F,
"location": F,
"start_date":F,
"end_date":F ,
"achievement": [F
]
}
],
"professional_summary": G,
"languages": [
{
"language": H,
"proficiency": H
}
],
"certifications": [
{
"name": I,
"issuing_organization": I,
"issuing_year": I
}
]
}"
"""
